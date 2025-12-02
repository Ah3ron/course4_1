#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Скрипт для преобразования Markdown файла в Word документ
с форматированием согласно требованиям к пояснительной записке.
"""

import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Константы форматирования
FONT_NAME = 'Times New Roman'
FONT_SIZE_NORMAL = Pt(14)
FONT_SIZE_HEADING_1 = Pt(16)
FONT_SIZE_HEADING_2 = Pt(14)
FONT_SIZE_HEADING_3 = Pt(14)
FONT_SIZE_TABLE = Pt(14)
FONT_SIZE_FIGURE = Pt(12)
FONT_COLOR_BLACK = RGBColor(0, 0, 0)


def setup_toc_styles(doc):
    """Настройка стилей TOC без отступов для всех уровней."""
    styles = doc.styles
    
    # Настраиваем стили TOC 1, TOC 2, TOC 3
    for toc_level in [1, 2, 3]:
        style_name = f'TOC {toc_level}'
        try:
            if style_name in [s.name for s in styles]:
                toc_style = styles[style_name]
            else:
                # Если стиль не существует, пропускаем
                continue
            
            # Убираем все отступы (левый, первый строки, висячий)
            toc_para = toc_style.paragraph_format
            toc_para.left_indent = Cm(0)
            toc_para.first_line_indent = Cm(0)
            toc_para.right_indent = Cm(0)
            
            # Устанавливаем шрифт Times New Roman
            toc_font = toc_style.font
            toc_font.name = FONT_NAME
            toc_font.size = FONT_SIZE_HEADING_2
            toc_font.color.rgb = FONT_COLOR_BLACK
        except Exception:
            # Если не удалось настроить стиль, пропускаем
            pass


def add_table_of_contents(doc):
    """Добавление автоматического оглавления в документ."""
    # Добавляем заголовок "СОДЕРЖАНИЕ"
    toc_heading = doc.add_paragraph('СОДЕРЖАНИЕ', style='Heading 1')
    
    # Добавляем пустую строку
    doc.add_paragraph()
    
    # Создаем параграф для поля TOC
    toc_para = doc.add_paragraph()
    
    # Создаем поле TOC (Table of Contents)
    # Получаем элемент параграфа
    p = toc_para._p
    
    # Создаем первый run для начала поля и инструкции
    r1 = OxmlElement('w:r')
    
    # Начало поля
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r1.append(fldChar_begin)
    
    # Инструкция для TOC
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # TOC с уровнями 1-3 (стили Heading 1, Heading 2, Heading 3), с гиперссылками, с точками
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r1.append(instrText)
    
    p.append(r1)
    
    # Создаем второй run для разделителя
    r2 = OxmlElement('w:r')
    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    r2.append(fldChar_separate)
    p.append(r2)
    
    # Создаем третий run для placeholder текста
    r3 = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = 'Для обновления оглавления щелкните правой кнопкой мыши и выберите "Обновить поле"'
    r3.append(t)
    p.append(r3)
    
    # Создаем четвертый run для конца поля
    r4 = OxmlElement('w:r')
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r4.append(fldChar_end)
    p.append(r4)
    
    # Настраиваем стили TOC без отступов
    setup_toc_styles(doc)
    
    # Добавляем разрыв страницы после оглавления
    doc.add_page_break()


def add_page_numbering(doc, start_from_intro=False):
    """Добавление нумерации страниц в верхний колонтитул по центру."""
    for section in doc.sections:
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем поле номера страницы
        run = header_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = 'PAGE'
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        
        run._element.append(fldChar1)
        run._element.append(instrText)
        run._element.append(fldChar2)
        
        # Форматирование номера страницы
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE_TABLE
        run.font.color.rgb = FONT_COLOR_BLACK


def setup_page_settings(doc):
    """Настройка параметров страницы согласно требованиям."""
    section = doc.sections[0]
    
    # Размер страницы А4
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    
    # Поля
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.5)
    
    # Расстояние от края до колонтитулов
    section.header_distance = Cm(2.5)
    section.footer_distance = Cm(2.0)


def setup_styles(doc):
    """Настройка стилей документа."""
    styles = doc.styles
    
    # Стиль для основного текста
    if 'Normal' in [s.name for s in styles]:
        normal_style = styles['Normal']
    else:
        normal_style = styles.add_style('Normal', 1)
    
    normal_font = normal_style.font
    normal_font.name = FONT_NAME
    normal_font.size = FONT_SIZE_NORMAL
    normal_font.color.rgb = FONT_COLOR_BLACK
    
    normal_para = normal_style.paragraph_format
    normal_para.first_line_indent = Cm(1.25)
    normal_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_para.line_spacing = 1.0
    normal_para.space_before = Pt(0)
    normal_para.space_after = Pt(0)
    
    # Стиль для ЗАГОЛОВОК 1
    if 'Heading 1' in [s.name for s in styles]:
        h1_style = styles['Heading 1']
    else:
        h1_style = styles.add_style('Heading 1', 1)
    
    h1_font = h1_style.font
    h1_font.name = FONT_NAME
    h1_font.size = FONT_SIZE_HEADING_1
    h1_font.bold = True
    h1_font.all_caps = True
    h1_font.color.rgb = FONT_COLOR_BLACK
    
    h1_para = h1_style.paragraph_format
    h1_para.first_line_indent = Cm(0)
    h1_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1_para.line_spacing = 1.0
    h1_para.space_before = Pt(6)
    h1_para.space_after = Pt(9)
    h1_para.keep_together = True
    h1_para.keep_with_next = True
    h1_para.widow_control = True
    
    # Стиль для Заголовок 2 (главы) - выровнен по центру
    if 'Heading 2' in [s.name for s in styles]:
        h2_style = styles['Heading 2']
    else:
        h2_style = styles.add_style('Heading 2', 1)
    
    h2_font = h2_style.font
    h2_font.name = FONT_NAME
    h2_font.size = FONT_SIZE_HEADING_2
    h2_font.bold = True
    h2_font.color.rgb = FONT_COLOR_BLACK
    
    h2_para = h2_style.paragraph_format
    h2_para.first_line_indent = Cm(0)  # Без отступа, так как по центру
    h2_para.alignment = WD_ALIGN_PARAGRAPH.CENTER  # Выравнивание по центру
    h2_para.line_spacing = 1.0
    h2_para.space_before = Pt(9)
    h2_para.space_after = Pt(6)
    h2_para.keep_together = True
    h2_para.keep_with_next = True
    h2_para.widow_control = True
    
    # Стиль для Заголовок 3
    if 'Heading 3' in [s.name for s in styles]:
        h3_style = styles['Heading 3']
    else:
        h3_style = styles.add_style('Heading 3', 1)
    
    h3_font = h3_style.font
    h3_font.name = FONT_NAME
    h3_font.size = FONT_SIZE_HEADING_3
    h3_font.bold = True
    h3_font.color.rgb = FONT_COLOR_BLACK
    
    h3_para = h3_style.paragraph_format
    h3_para.first_line_indent = Cm(1.25)
    h3_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h3_para.line_spacing = 1.0
    h3_para.space_before = Pt(9)
    h3_para.space_after = Pt(6)
    h3_para.keep_together = True
    h3_para.keep_with_next = True
    h3_para.widow_control = True


def parse_markdown_formatting(text, para):
    """Парсинг markdown форматирования (жирный текст) и добавление в параграф."""
    if not text:
        return
    
    # Разбиваем текст на части, сохраняя разделители
    # Обрабатываем жирный текст **текст** или __текст__
    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__)', text)
    
    for part in parts:
        if not part:
            continue
        
        # Жирный текст **текст** или __текст__
        if re.match(r'^\*\*[^*]+\*\*$', part):
            bold_text = part[2:-2]
            run = para.add_run(bold_text)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.color.rgb = FONT_COLOR_BLACK
        elif re.match(r'^__[^_]+__$', part):
            bold_text = part[2:-2]
            run = para.add_run(bold_text)
            run.bold = True
            run.font.name = FONT_NAME
            run.font.color.rgb = FONT_COLOR_BLACK
        else:
            # Обычный текст
            run = para.add_run(part)
            run.font.name = FONT_NAME
            run.font.color.rgb = FONT_COLOR_BLACK


def parse_markdown_table(table_text):
    """Парсинг markdown таблицы."""
    lines = [line.strip() for line in table_text.strip().split('\n') if line.strip()]
    if not lines:
        return []
    
    # Пропускаем разделитель (вторую строку)
    data_lines = [line for line in lines if not re.match(r'^[\|\s\-\:]+$', line)]
    
    rows = []
    for line in data_lines:
        # Убираем начальные и конечные |
        cells = [cell.strip() for cell in line.split('|') if cell.strip()]
        if cells:
            rows.append(cells)
    
    return rows




def get_or_create_numbering(doc):
    """Получает или создает numbering part для документа.
    
    Возвращает корневой элемент numbering.
    """
    part = doc.part
    
    # Пытаемся получить существующий numbering_part
    try:
        numbering_part = part.numbering_part
        if numbering_part is not None:
            # Возвращаем элемент numbering
            return numbering_part.element
    except AttributeError:
        pass
    
    # Если numbering_part не существует, создаем его
    numbering_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?><w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"></w:numbering>'
    from docx.oxml import parse_xml
    numbering_element = parse_xml(numbering_xml)
    
    # Создаем numbering part через package
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    package = part.package
    
    # Создаем новую часть
    numbering_part = package.get_or_add_part(RT.NUMBERING, 'word/numbering.xml')
    numbering_part.element = numbering_element
    
    # Связываем с основным документом
    part.relate_to(numbering_part, RT.NUMBERING)
    
    return numbering_element


def create_numbering_definitions(doc):
    """Создание определений нумерации для нумерованных и маркированных списков.
    
    Возвращает словарь с информацией о созданных определениях.
    """
    # Получаем или создаем numbering
    numbering = get_or_create_numbering(doc)
    
    # Создаем абстрактные определения для нумерованных списков
    # Используем abstractNumId начиная с 0
    abstract_num_id = 0
    
    # Абстрактное определение для нумерованных списков
    abstract_num = OxmlElement('w:abstractNum')
    abstract_num.set(qn('w:abstractNumId'), str(abstract_num_id))
    
    # Для каждого уровня создаем формат нумерации
    for level in range(9):  # Поддерживаем до 9 уровней
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))
        
        # Формат нумерации зависит от уровня
        if level == 0:
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), 'decimal')  # 1, 2, 3...
            lvl.append(numFmt)
        elif level == 1:
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), 'lowerLetter')  # a, b, c...
            lvl.append(numFmt)
        elif level == 2:
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), 'lowerRoman')  # i, ii, iii...
            lvl.append(numFmt)
        elif level == 3:
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), 'decimal')  # 1, 2, 3...
            lvl.append(numFmt)
        else:
            # Для остальных уровней используем decimal
            numFmt = OxmlElement('w:numFmt')
            numFmt.set(qn('w:val'), 'decimal')
            lvl.append(numFmt)
        
        # Начало нумерации
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        
        # Формат текста для нумерации
        lvlText = OxmlElement('w:lvlText')
        if level == 0:
            lvlText.set(qn('w:val'), '%1.')
        elif level == 1:
            lvlText.set(qn('w:val'), '%2.')
        elif level == 2:
            lvlText.set(qn('w:val'), '%3.')
        else:
            lvlText.set(qn('w:val'), f'%{level + 1}.')
        lvl.append(lvlText)
        
        # Выравнивание
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        
        # Отступы (будут переопределены в setup_list_formatting)
        pPr = OxmlElement('w:pPr')
        indent = OxmlElement('w:ind')
        indent.set(qn('w:left'), '720')  # Базовый отступ
        indent.set(qn('w:hanging'), '360')  # Висячий отступ
        pPr.append(indent)
        lvl.append(pPr)
        
        abstract_num.append(lvl)
    
    numbering.append(abstract_num)
    
    # Создаем абстрактное определение для маркированных списков
    abstract_num_id_bullet = 1
    abstract_num_bullet = OxmlElement('w:abstractNum')
    abstract_num_bullet.set(qn('w:abstractNumId'), str(abstract_num_id_bullet))
    
    for level in range(9):
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), str(level))
        
        # Формат маркера
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'bullet')
        lvl.append(numFmt)
        
        # Маркер
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), '•')
        lvl.append(lvlText)
        
        # Выравнивание
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)
        
        # Отступы
        pPr = OxmlElement('w:pPr')
        indent = OxmlElement('w:ind')
        indent.set(qn('w:left'), '720')
        indent.set(qn('w:hanging'), '360')
        pPr.append(indent)
        lvl.append(pPr)
        
        abstract_num_bullet.append(lvl)
    
    numbering.append(abstract_num_bullet)
    
    return {
        'numbered_abstract_id': abstract_num_id,
        'bullet_abstract_id': abstract_num_id_bullet,
        'next_num_id': 1  # Следующий доступный numId
    }


def create_num_instance(numbering, abstract_num_id, num_id):
    """Создание экземпляра нумерации (num) на основе абстрактного определения."""
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), str(num_id))
    
    abstractNumId = OxmlElement('w:abstractNumId')
    abstractNumId.set(qn('w:val'), str(abstract_num_id))
    num.append(abstractNumId)
    
    numbering.append(num)


def setup_list_formatting(para, is_numbered, level, num_id):
    """Настройка форматирования списка через XML для автоматической нумерации/маркировки.
    
    Args:
        para: Параграф для форматирования
        is_numbered: True для нумерованного списка, False для маркированного
        level: Уровень вложенности (0-based)
        num_id: ID нумерации (уникальный для каждого списка)
    """
    p = para._p
    pPr = p.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p.insert(0, pPr)
    
    # Удаляем существующую нумерацию, если есть
    existing_numPr = pPr.find(qn('w:numPr'))
    if existing_numPr is not None:
        pPr.remove(existing_numPr)
    
    # Создаем элемент нумерации
    numPr = OxmlElement('w:numPr')
    
    # Уровень вложенности (0-based)
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    
    # ID нумерации
    numId = OxmlElement('w:numId')
    numId.set(qn('w:val'), str(num_id))
    numPr.append(numId)
    
    pPr.append(numPr)


def format_table_borders(table):
    """Установка черных границ и прозрачного фона для таблицы."""
    # Устанавливаем границы для всей таблицы
    tbl = table._tbl
    
    # Получаем или создаем tblPr через прямое обращение к XML
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        # Создаем tblPr элемент
        tblPr = OxmlElement('w:tblPr')
        # Вставляем в начало таблицы
        if len(tbl) > 0:
            tbl.insert(0, tblPr)
        else:
            tbl.append(tblPr)
    
    # Удаляем существующие границы, если есть
    existing_borders = tblPr.find(qn('w:tblBorders'))
    if existing_borders is not None:
        tblPr.remove(existing_borders)
    
    # Создаем новые границы
    tblBorders = OxmlElement('w:tblBorders')
    
    # Черные границы для всех сторон
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Размер границы
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Черный цвет
        tblBorders.append(border)
    
    tblPr.append(tblBorders)
    
    # Устанавливаем прозрачный фон для всех ячеек
    for row in table.rows:
        for cell in row.cells:
            # Убираем заливку ячейки
            tc = cell._tc
            
            # Получаем или создаем tcPr
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            
            # Удаляем существующую заливку, если есть
            existing_shd = tcPr.find(qn('w:shd'))
            if existing_shd is not None:
                tcPr.remove(existing_shd)
            
            # Создаем новую заливку с прозрачным фоном
            cell_shading = OxmlElement('w:shd')
            cell_shading.set(qn('w:val'), 'clear')
            cell_shading.set(qn('w:fill'), 'auto')
            tcPr.append(cell_shading)


def add_table_to_doc(doc, table_data, section_num, table_num, table_name=None):
    """Добавление таблицы в документ с правильным форматированием."""
    if not table_data or len(table_data) < 1:
        return
    
    # Добавляем заголовок таблицы (подпись над таблицей)
    table_caption = doc.add_paragraph()
    table_caption.alignment = WD_ALIGN_PARAGRAPH.LEFT
    table_caption.paragraph_format.first_line_indent = Cm(1.25)
    table_caption.paragraph_format.space_before = Pt(6)
    table_caption.paragraph_format.space_after = Pt(6)
    
    run = table_caption.add_run(f'Таблица {section_num}.{table_num} – ')
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE_NORMAL
    run.bold = True
    run.font.color.rgb = FONT_COLOR_BLACK
    
    # Название таблицы
    if table_name:
        caption_text = table_name
    elif len(table_data) > 0 and len(table_data[0]) > 0:
        # Пытаемся использовать первую ячейку первой строки как название
        caption_text = table_data[0][0] if table_data[0] else f'Таблица {section_num}.{table_num}'
    else:
        caption_text = f'Таблица {section_num}.{table_num}'
    
    caption_run = table_caption.add_run(caption_text)
    caption_run.font.name = FONT_NAME
    caption_run.font.size = FONT_SIZE_NORMAL
    caption_run.bold = True
    caption_run.font.color.rgb = FONT_COLOR_BLACK
    
    # Определяем количество строк и столбцов
    # Первая строка - заголовок таблицы, остальные - данные
    header_row = table_data[0] if table_data else []
    data_rows = table_data[1:] if len(table_data) > 1 else []
    
    # Определяем максимальное количество столбцов
    max_cols = max(len(row) for row in table_data) if table_data else len(header_row)
    if max_cols == 0:
        return
    
    # Создаем таблицу: 1 строка заголовка + строки данных
    table_rows = 1 + len(data_rows)
    table = doc.add_table(rows=table_rows, cols=max_cols)
    # Не используем стиль, чтобы иметь полный контроль над форматированием
    table.style = None
    
    # Добавляем заголовок таблицы (первая строка) - жирным шрифтом
    header_cells = table.rows[0].cells
    for j, cell_data in enumerate(header_row):
        if j < len(header_cells):
            cell = header_cells[j]
            cell.text = ''  # Очищаем ячейку
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)
            # Обрабатываем markdown форматирование в заголовке
            parse_markdown_formatting(cell_data, para)
            # Устанавливаем шрифт для всех runs в заголовке - жирным
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_TABLE
                run.font.color.rgb = FONT_COLOR_BLACK
                run.bold = True  # Заголовок жирным
    
    # Заполняем пустые ячейки заголовка, если столбцов больше чем данных
    for j in range(len(header_row), max_cols):
        if j < len(header_cells):
            cell = header_cells[j]
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            para.paragraph_format.line_spacing = Pt(18)
    
    # Добавляем строки данных
    for i, row_data in enumerate(data_rows):
        row_index = i + 1  # +1 потому что первая строка - заголовок
        if row_index >= len(table.rows):
            break
        data_cells = table.rows[row_index].cells
        for j, cell_data in enumerate(row_data):
            if j < len(data_cells):
                cell = data_cells[j]
                cell.text = ''  # Очищаем ячейку
                para = cell.paragraphs[0]
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(18)
                # Обрабатываем markdown форматирование в ячейке
                parse_markdown_formatting(cell_data, para)
                # Устанавливаем шрифт для всех runs в ячейке - обычный (не жирный)
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_TABLE
                    run.font.color.rgb = FONT_COLOR_BLACK
                    run.bold = False  # Данные не жирным
        
        # Заполняем пустые ячейки, если столбцов больше чем данных
        for j in range(len(row_data), max_cols):
            if j < len(data_cells):
                cell = data_cells[j]
                cell.text = ''
                para = cell.paragraphs[0]
                para.paragraph_format.first_line_indent = Cm(0)
                para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(18)
    
    # Применяем черные границы и прозрачный фон
    format_table_borders(table)
    
    # Добавляем отступ после таблицы
    doc.add_paragraph()


def process_markdown_file(md_file_path, output_path):
    """Основная функция обработки markdown файла."""
    with open(md_file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    setup_page_settings(doc)
    setup_styles(doc)
    
    # Инициализируем определения нумерации
    numbering_info = create_numbering_definitions(doc)
    
    # Добавляем нумерацию страниц (будет применяться ко всем разделам)
    add_page_numbering(doc)
    
    lines = content.split('\n')
    i = 0
    current_section = 0
    current_subsection = 0
    current_point = 0
    table_num = 0
    figure_num = 0
    in_table = False
    table_lines = []
    current_list_level = 0
    toc_inserted = False
    introduction_passed = False  # Флаг для отслеживания, прошло ли введение
    
    # Счетчики для нумерации таблиц и рисунков по разделам
    section_table_counters = {}
    section_figure_counters = {}
    
    # Отслеживание состояния списков
    list_context = {
        'in_list': False,  # Находимся ли мы внутри списка
        'list_type': None,  # 'numbered' или 'bullet' или None
        'list_level': None,  # Текущий уровень вложенности
        'list_num_id': None,  # Текущий numId (уникальный для каждого списка)
        'list_levels': {}  # Словарь для отслеживания numId на каждом уровне вложенности
    }
    
    # Счетчик для создания уникальных numId
    next_num_id = numbering_info['next_num_id']
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Обработка заголовков
        if line.startswith('#'):
            # Закрываем таблицу, если она была открыта
            if in_table:
                table_data = parse_markdown_table('\n'.join(table_lines))
                if table_data:
                    section_num = current_section if current_section > 0 else 1
                    if section_num not in section_table_counters:
                        section_table_counters[section_num] = 0
                    section_table_counters[section_num] += 1
                    add_table_to_doc(doc, table_data, section_num, section_table_counters[section_num])
                in_table = False
                table_lines = []
            
            # Заголовок прерывает список
            list_context['in_list'] = False
            list_context['list_type'] = None
            list_context['list_level'] = None
            list_context['list_num_id'] = None
            list_context['list_levels'] = {}
            
            # Определяем уровень заголовка
            level = len(line) - len(line.lstrip('#'))
            header_text = line.lstrip('#').strip()
            
            # Удаляем разделители (---)
            if header_text == '' or header_text == '-' or header_text.startswith('---'):
                i += 1
                continue
            
            # Проверка на приложение
            is_appendix = 'ПРИЛОЖЕНИЕ' in header_text.upper() or 'Приложение' in header_text
            
            # Проверка на введение
            is_introduction = 'ВВЕДЕНИЕ' in header_text.upper() or 'Введение' in header_text
            
            # Если это введение, устанавливаем флаг
            if is_introduction and level == 2:
                introduction_passed = True
            
            # Если это заголовок уровня 2 (главы)
            if level == 2:
                # Если это первый заголовок уровня 2, добавляем оглавление перед ним
                if not toc_inserted:
                    if doc.paragraphs:
                        doc.add_page_break()
                    add_table_of_contents(doc)
                    toc_inserted = True
                else:
                    # Для всех последующих заголовков уровня 2 добавляем разрыв страницы
                    doc.add_page_break()
                
                # Для обычных разделов увеличиваем счетчик
                if not is_appendix:
                    current_section += 1
                    current_subsection = 0
                    current_point = 0
            
            # Добавляем заголовок с правильным стилем
            if level == 1:
                para = doc.add_paragraph(style='Heading 1')
                # Извлекаем номер раздела из текста (если есть)
                if not is_appendix:
                    match = re.match(r'^(\d+)\s', header_text)
                    if match:
                        current_section = int(match.group(1))
            elif level == 2:
                para = doc.add_paragraph(style='Heading 2')
                # Извлекаем номер раздела (для глав) или подраздела
                # Сначала проверяем формат главы: "1 Название"
                match = re.match(r'^(\d+)\s', header_text)
                if match:
                    current_section = int(match.group(1))
                    current_subsection = 0
                    current_point = 0
                    # Если прошло введение, меняем формат на "ГЛАВА 1"
                    if introduction_passed and not is_appendix and not is_introduction:
                        # Заменяем "1" на "ГЛАВА 1" в тексте заголовка
                        header_text = re.sub(r'^(\d+)\s', r'ГЛАВА \1 ', header_text)
                else:
                    # Проверяем формат подраздела: "1.2 Название"
                    match = re.match(r'^(\d+)\.(\d+)\s', header_text)
                    if match:
                        current_section = int(match.group(1))
                        current_subsection = int(match.group(2))
            elif level == 3:
                para = doc.add_paragraph(style='Heading 3')
                # Извлекаем номер пункта
                match = re.match(r'^(\d+)\.(\d+)\.(\d+)\s', header_text)
                if match:
                    current_section = int(match.group(1))
                    current_subsection = int(match.group(2))
                    current_point = int(match.group(3))
            elif level == 4:
                para = doc.add_paragraph(style='Heading 3')  # Заголовок 4 уровня тоже как Heading 3
            
            # Обрабатываем форматирование в заголовке (если есть)
            parse_markdown_formatting(header_text, para)
            # Устанавливаем шрифт для всех runs в параграфе (переопределяем стиль)
            for run in para.runs:
                run.font.name = FONT_NAME
                if level == 1:
                    run.font.size = FONT_SIZE_HEADING_1
                    run.bold = True
                    run.font.all_caps = True
                else:
                    run.font.size = FONT_SIZE_HEADING_2
                    run.bold = True
            
            i += 1
            continue
        
        # Обработка таблиц
        if '|' in line:
            # Проверяем, не является ли это разделителем таблицы
            if re.match(r'^[\|\s\-\:]+$', line.strip()):
                # Это разделитель, добавляем в таблицу
                if in_table:
                    table_lines.append(line)
            elif not line.strip().startswith('---'):
                if not in_table:
                    in_table = True
                    table_lines = []
                table_lines.append(line)
            i += 1
            continue
        elif in_table:
            # Завершаем таблицу, если следующая строка не является частью таблицы
            table_data = parse_markdown_table('\n'.join(table_lines))
            if table_data:
                section_num = current_section if current_section > 0 else 1
                if section_num not in section_table_counters:
                    section_table_counters[section_num] = 0
                section_table_counters[section_num] += 1
                add_table_to_doc(doc, table_data, section_num, section_table_counters[section_num])
            in_table = False
            table_lines = []
            # Таблица прерывает список
            list_context['in_list'] = False
            list_context['list_type'] = None
            list_context['list_level'] = None
            list_context['list_num_id'] = None
            list_context['list_levels'] = {}
        
        # Обработка изображений
        if re.match(r'^!\[.*?\]\(.*?\)', line):
            # Извлекаем название рисунка
            match = re.match(r'!\[(.*?)\]\(.*?\)', line)
            if match:
                figure_name = match.group(1)
                section_num = current_section if current_section > 0 else 1
                if section_num not in section_figure_counters:
                    section_figure_counters[section_num] = 0
                section_figure_counters[section_num] += 1
                
                # Добавляем подпись под рисунком
                fig_para = doc.add_paragraph()
                fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                fig_run = fig_para.add_run(f'Рисунок {section_num}.{section_figure_counters[section_num]} – {figure_name}')
                fig_run.font.name = FONT_NAME
                fig_run.font.size = FONT_SIZE_FIGURE
                fig_run.font.color.rgb = FONT_COLOR_BLACK
                fig_para.paragraph_format.line_spacing = 1.0
                fig_para.paragraph_format.space_before = Pt(6)
                fig_para.paragraph_format.space_after = Pt(6)
            
            i += 1
            continue
        
        # Обработка списков
        is_bullet_list = re.match(r'^\s*[-*+]\s+', line)
        is_numbered_list = re.match(r'^\s*\d+\.\s+', line)
        
        if is_bullet_list or is_numbered_list:
            # Определяем уровень вложенности по отступам
            # В markdown вложенные списки обычно имеют отступ 2-4 пробела на уровень
            leading_spaces = len(line) - len(line.lstrip())
            # Определяем уровень: 0-2 пробела = уровень 0, 3-5 = уровень 1, 6-8 = уровень 2 и т.д.
            indent_level = leading_spaces // 3 if leading_spaces > 2 else 0
            # Ограничиваем уровень до 8 (максимум для стилей Word)
            indent_level = min(indent_level, 8)
            
            # Определяем тип текущего списка
            current_list_type = 'numbered' if is_numbered_list else 'bullet'
            
            # Определяем, начинается ли новый список
            # Новый список начинается, если:
            # 1. Мы не находимся в списке (in_list = False)
            # 2. Тип списка изменился (numbered <-> bullet)
            # 3. Уровень вложенности уменьшился, и мы вернулись к уровню, для которого нет сохраненного numId
            is_new_list = False
            
            if not list_context['in_list']:
                # Начинаем новый список
                is_new_list = True
            elif list_context['list_type'] != current_list_type:
                # Тип списка изменился - начинаем новый список
                is_new_list = True
            elif indent_level < list_context['list_level']:
                # Уровень уменьшился - проверяем, есть ли сохраненный numId для этого уровня
                if indent_level in list_context['list_levels']:
                    # Есть сохраненный numId - продолжаем существующий список на этом уровне
                    is_new_list = False
                else:
                    # Нет сохраненного numId - это новый список
                    is_new_list = True
            elif indent_level > list_context['list_level']:
                # Уровень увеличился - это вложенный список, используем тот же numId
                is_new_list = False
            else:
                # Уровень не изменился - продолжаем тот же список
                is_new_list = False
            
            # Определяем numId для текущего элемента списка
            if is_new_list:
                # Создаем новый список - нужен новый numId
                num_id = next_num_id
                next_num_id += 1
                
                # Создаем экземпляр нумерации в документе
                numbering = get_or_create_numbering(doc)
                if current_list_type == 'numbered':
                    abstract_num_id = numbering_info['numbered_abstract_id']
                else:
                    abstract_num_id = numbering_info['bullet_abstract_id']
                create_num_instance(numbering, abstract_num_id, num_id)
                
                # Сохраняем numId для текущего уровня
                list_context['list_levels'][indent_level] = num_id
            else:
                # Продолжаем существующий список
                # Если есть сохраненный numId для текущего уровня, используем его
                if indent_level in list_context['list_levels']:
                    num_id = list_context['list_levels'][indent_level]
                elif list_context['list_num_id'] is not None:
                    # Используем numId из контекста (для вложенных уровней)
                    num_id = list_context['list_num_id']
                    # Сохраняем для текущего уровня
                    list_context['list_levels'][indent_level] = num_id
                else:
                    # Если numId нет в контексте, создаем новый (не должно происходить в нормальных условиях)
                    num_id = next_num_id
                    next_num_id += 1
                    numbering = get_or_create_numbering(doc)
                    if current_list_type == 'numbered':
                        abstract_num_id = numbering_info['numbered_abstract_id']
                    else:
                        abstract_num_id = numbering_info['bullet_abstract_id']
                    create_num_instance(numbering, abstract_num_id, num_id)
                    list_context['list_levels'][indent_level] = num_id
            
            # Извлекаем текст элемента списка (без маркера/номера)
            if is_numbered_list:
                list_text = re.sub(r'^\s*\d+\.\s+', '', line)
            else:
                list_text = re.sub(r'^\s*[-*+]\s+', '', line)
            
            # Создаем параграф с базовым стилем и добавляем нумерацию через XML
            para = doc.add_paragraph(style='Normal')
            # Настраиваем нумерацию/маркировку через XML
            setup_list_formatting(para, is_numbered_list, indent_level, num_id)
            
            # Настройка отступов для списка
            # Базовый отступ для списков (как у обычного текста)
            base_indent = Cm(1.25)
            indent_per_level = Cm(0.5)
            
            # Для первого уровня: отступ слева 1.25 см, висячий отступ для маркера/номера
            # Для вложенных уровней: увеличиваем отступ слева
            para.paragraph_format.left_indent = base_indent + (indent_per_level * indent_level)
            para.paragraph_format.first_line_indent = Cm(-0.5)  # Висячий отступ для маркера/номера
            
            # Обрабатываем форматирование в тексте списка
            parse_markdown_formatting(list_text, para)
            
            # Устанавливаем шрифт для всех runs в параграфе
            for run in para.runs:
                run.font.name = FONT_NAME
                run.font.size = FONT_SIZE_NORMAL
                run.font.color.rgb = FONT_COLOR_BLACK
            
            # Обновляем контекст списка
            list_context['in_list'] = True
            list_context['list_type'] = current_list_type
            list_context['list_level'] = indent_level
            list_context['list_num_id'] = num_id
            # Сохраняем numId для текущего уровня (если еще не сохранен)
            if indent_level not in list_context['list_levels']:
                list_context['list_levels'][indent_level] = num_id
            
            i += 1
            continue
        
        # Пропускаем пустые строки и разделители
        if not line.strip() or line.strip() == '---':
            # Пустая строка прерывает список, если следующая строка не является продолжением списка
            # Проверяем следующую строку
            if i + 1 < len(lines):
                next_line = lines[i + 1].rstrip()
                next_is_list = bool(re.match(r'^\s*[-*+]\s+', next_line) or re.match(r'^\s*\d+\.\s+', next_line))
                if not next_is_list:
                    # Следующая строка не список - прерываем текущий список
                    list_context['in_list'] = False
                    list_context['list_type'] = None
                    list_context['list_level'] = None
                    list_context['list_num_id'] = None
                    list_context['list_levels'] = {}
            else:
                # Это последняя строка - прерываем список
                list_context['in_list'] = False
                list_context['list_type'] = None
                list_context['list_level'] = None
                list_context['list_num_id'] = None
                list_context['list_levels'] = {}
            i += 1
            continue
        
        # Обработка обычного текста
        if line.strip():
            # Обычный текст прерывает список
            list_context['in_list'] = False
            list_context['list_type'] = None
            list_context['list_level'] = None
            list_context['list_num_id'] = None
            list_context['list_levels'] = {}
            
            # Проверяем, не является ли это формулой (упрощенная проверка)
            is_formula = (re.search(r'[XZ]\s*=\s*[\d\.]+\s*[X\+\-]', line) or 
                         re.search(r'[XZ]\s*=\s*[\d\.]+\s*X', line) or
                         'Где:' in line or 'где:' in line or
                         re.match(r'^[XZ]\s*=\s*', line.strip()))
            
            if is_formula:
                # Формула - добавляем с отступами
                para = doc.add_paragraph(style='Normal')
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                # Для формул убираем абзацный отступ
                para.paragraph_format.first_line_indent = Cm(0)
                # Обрабатываем форматирование в формуле
                parse_markdown_formatting(line.strip(), para)
                # Устанавливаем шрифт для всех runs в параграфе
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_NORMAL
            else:
                # Обычный текст - обрабатываем markdown форматирование
                para = doc.add_paragraph(style='Normal')
                parse_markdown_formatting(line.strip(), para)
                # Устанавливаем шрифт для всех runs в параграфе
                for run in para.runs:
                    run.font.name = FONT_NAME
                    run.font.size = FONT_SIZE_NORMAL
        
        i += 1
    
    # Обработка последней таблицы, если она была открыта
    if in_table:
        table_data = parse_markdown_table('\n'.join(table_lines))
        if table_data:
            section_num = current_section if current_section > 0 else 1
            if section_num not in section_table_counters:
                section_table_counters[section_num] = 0
            section_table_counters[section_num] += 1
            add_table_to_doc(doc, table_data, section_num, section_table_counters[section_num])
    
    # Сохраняем документ
    doc.save(output_path)
    print(f'Документ успешно создан: {output_path}')


def main():
    """Главная функция."""
    if len(sys.argv) < 2:
        print("Использование: python md_to_word.py <путь_к_md_файлу> [путь_к_output_файлу]")
        print("Пример: python md_to_word.py Пояснительная_записка.md Пояснительная_записка.docx")
        sys.exit(1)
    
    md_file = Path(sys.argv[1])
    if not md_file.exists():
        print(f"Ошибка: файл {md_file} не найден!")
        sys.exit(1)
    
    if len(sys.argv) >= 3:
        output_file = Path(sys.argv[2])
    else:
        output_file = md_file.with_suffix('.docx')
    
    try:
        process_markdown_file(md_file, output_file)
    except Exception as e:
        print(f"Ошибка при обработке файла: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == '__main__':
    main()

